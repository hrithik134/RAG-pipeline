"""
Text extraction service for different document formats.

This module provides text extraction from PDF, DOCX, and TXT files with
page count detection and fallback mechanisms.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import chardet
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from pdfminer.high_level import extract_text as pdfminer_extract_text

from app.config import settings
from app.utils.exceptions import ExtractionError, PageLimitExceededError


@dataclass
class ExtractedText:
    """Container for extracted text and metadata."""
    
    text: str
    page_count: int
    char_count: int
    metadata: dict


class BaseExtractor(ABC):
    """Abstract base class for text extractors."""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> ExtractedText:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractedText object with text and metadata
            
        Raises:
            ExtractionError: If extraction fails
        """
        pass
    
    @abstractmethod
    def get_page_count(self, file_path: str) -> int:
        """
        Get the page count of a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Number of pages
            
        Raises:
            ExtractionError: If page count cannot be determined
        """
        pass
    
    def validate_page_count(self, page_count: int, filename: str) -> None:
        """
        Validate that page count is within limits.
        
        Args:
            page_count: Number of pages
            filename: Name of the file
            
        Raises:
            PageLimitExceededError: If page count exceeds limit
        """
        max_pages = settings.max_pages_per_document
        
        if page_count > max_pages:
            raise PageLimitExceededError(
                filename=filename,
                page_count=page_count,
                max_pages=max_pages
            )


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using PyMuPDF with pdfminer fallback."""
    
    def get_page_count(self, file_path: str) -> int:
        """Get page count from PDF."""
        try:
            with fitz.open(file_path) as doc:
                return len(doc)
        except Exception as e:
            raise ExtractionError(
                filename=Path(file_path).name,
                file_type="pdf",
                error_details=f"Failed to get page count: {str(e)}"
            )
    
    def extract_text(self, file_path: str) -> ExtractedText:
        """
        Extract text from PDF using PyMuPDF, fallback to pdfminer.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ExtractedText object
            
        Raises:
            ExtractionError: If extraction fails
            PageLimitExceededError: If page count exceeds limit
        """
        filename = Path(file_path).name
        
        try:
            # Try PyMuPDF first (faster and more accurate)
            return self._extract_with_pymupdf(file_path, filename)
        except Exception as pymupdf_error:
            # Fallback to pdfminer.six
            try:
                return self._extract_with_pdfminer(file_path, filename)
            except Exception as pdfminer_error:
                raise ExtractionError(
                    filename=filename,
                    file_type="pdf",
                    error_details=(
                        f"PyMuPDF failed: {str(pymupdf_error)}; "
                        f"pdfminer failed: {str(pdfminer_error)}"
                    )
                )
    
    def _extract_with_pymupdf(self, file_path: str, filename: str) -> ExtractedText:
        """Extract text using PyMuPDF."""
        with fitz.open(file_path) as doc:
            page_count = len(doc)
            
            # Validate page count
            self.validate_page_count(page_count, filename)
            
            # Extract text from all pages
            text_parts = []
            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            
            return ExtractedText(
                text=full_text,
                page_count=page_count,
                char_count=len(full_text),
                metadata={
                    "extractor": "pymupdf",
                    "pdf_metadata": doc.metadata
                }
            )
    
    def _extract_with_pdfminer(self, file_path: str, filename: str) -> ExtractedText:
        """Extract text using pdfminer.six as fallback."""
        # Get page count first
        page_count = self.get_page_count(file_path)
        
        # Validate page count
        self.validate_page_count(page_count, filename)
        
        # Extract text
        text = pdfminer_extract_text(file_path)
        
        return ExtractedText(
            text=text,
            page_count=page_count,
            char_count=len(text),
            metadata={
                "extractor": "pdfminer"
            }
        )


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""
    
    def get_page_count(self, file_path: str) -> int:
        """
        Estimate page count for DOCX.
        
        DOCX doesn't have explicit pages, so we estimate based on character count.
        Assumes ~1800 characters per page (standard formatting).
        """
        try:
            doc = DocxDocument(file_path)
            
            # Count total characters
            total_chars = 0
            for paragraph in doc.paragraphs:
                total_chars += len(paragraph.text)
            
            # Estimate pages (1800 chars per page is a reasonable estimate)
            estimated_pages = max(1, total_chars // 1800)
            
            return estimated_pages
            
        except Exception as e:
            raise ExtractionError(
                filename=Path(file_path).name,
                file_type="docx",
                error_details=f"Failed to get page count: {str(e)}"
            )
    
    def extract_text(self, file_path: str) -> ExtractedText:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ExtractedText object
            
        Raises:
            ExtractionError: If extraction fails
            PageLimitExceededError: If estimated page count exceeds limit
        """
        filename = Path(file_path).name
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            page_count = self.get_page_count(file_path)
            
            # Validate page count
            self.validate_page_count(page_count, filename)
            
            return ExtractedText(
                text=full_text,
                page_count=page_count,
                char_count=len(full_text),
                metadata={
                    "extractor": "python-docx",
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables)
                }
            )
            
        except PageLimitExceededError:
            raise
        except Exception as e:
            raise ExtractionError(
                filename=filename,
                file_type="docx",
                error_details=str(e)
            )


class TXTExtractor(BaseExtractor):
    """Extract text from TXT files with encoding detection."""
    
    def get_page_count(self, file_path: str) -> int:
        """TXT files are considered as 1 page."""
        return 1
    
    def extract_text(self, file_path: str) -> ExtractedText:
        """
        Extract text from TXT file with automatic encoding detection.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            ExtractedText object
            
        Raises:
            ExtractionError: If extraction fails
        """
        filename = Path(file_path).name
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
            
            return ExtractedText(
                text=text,
                page_count=1,
                char_count=len(text),
                metadata={
                    "extractor": "txt",
                    "encoding": encoding,
                    "confidence": result.get('confidence', 0)
                }
            )
            
        except Exception as e:
            raise ExtractionError(
                filename=filename,
                file_type="txt",
                error_details=str(e)
            )


class MarkdownExtractor(BaseExtractor):
    """Extract text from Markdown files."""
    
    def get_page_count(self, file_path: str) -> int:
        """Markdown files are considered as 1 page."""
        return 1
    
    def extract_text(self, file_path: str) -> ExtractedText:
        """
        Extract text from Markdown file.
        
        Args:
            file_path: Path to MD file
            
        Returns:
            ExtractedText object
            
        Raises:
            ExtractionError: If extraction fails
        """
        filename = Path(file_path).name
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
            
            return ExtractedText(
                text=text,
                page_count=1,
                char_count=len(text),
                metadata={
                    "extractor": "markdown",
                    "encoding": encoding,
                    "confidence": result.get('confidence', 0)
                }
            )
            
        except Exception as e:
            raise ExtractionError(
                filename=filename,
                file_type="md",
                error_details=str(e)
            )


class ExtractorFactory:
    """Factory for creating appropriate text extractors."""
    
    _extractors = {
        '.pdf': PDFExtractor,
        '.docx': DOCXExtractor,
        '.txt': TXTExtractor,
        '.md': MarkdownExtractor,
    }
    
    @classmethod
    def get_extractor(cls, file_path: str) -> BaseExtractor:
        """
        Get appropriate extractor for file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Instance of appropriate extractor
            
        Raises:
            ExtractionError: If file type is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        extractor_class = cls._extractors.get(file_ext)
        
        if not extractor_class:
            raise ExtractionError(
                filename=Path(file_path).name,
                file_type=file_ext,
                error_details=f"No extractor available for file type: {file_ext}"
            )
        
        return extractor_class()
    
    @classmethod
    def extract_text(cls, file_path: str) -> ExtractedText:
        """
        Convenience method to extract text using appropriate extractor.
        
        Args:
            file_path: Path to the file
            
        Returns:
            ExtractedText object
        """
        extractor = cls.get_extractor(file_path)
        return extractor.extract_text(file_path)

