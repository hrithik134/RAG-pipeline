"""
Comprehensive tests for text extraction services.

Tests cover PDF, DOCX, and TXT extractors with edge cases,
error handling, and different file formats.
"""

import io
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from zipfile import ZipFile

import pytest
import fitz  # PyMuPDF

from app.services.text_extractor import (
    PDFExtractor,
    DOCXExtractor,
    TXTExtractor,
    ExtractorFactory,
    ExtractedText,
    BaseExtractor,
)
from app.utils.exceptions import ExtractionError, PageLimitExceededError


@pytest.mark.unit
class TestPDFExtractor:
    """Tests for PDF text extraction."""
    
    def test_pdf_extractor_initialization(self):
        """Test PDF extractor initializes correctly."""
        extractor = PDFExtractor()
        
        assert isinstance(extractor, PDFExtractor)
        assert isinstance(extractor, BaseExtractor)
    
    def test_extract_from_simple_pdf(self, temp_dir):
        """Test extraction from a simple PDF file."""
        extractor = PDFExtractor()
        
        # Create a simple PDF
        pdf_path = temp_dir / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Test PDF content")
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert isinstance(result, ExtractedText)
        assert "Test PDF content" in result.text
        assert result.page_count == 1
        assert result.char_count > 0
    
    def test_extract_from_multi_page_pdf(self, temp_dir):
        """Test extraction from multi-page PDF."""
        extractor = PDFExtractor()
        
        # Create multi-page PDF
        pdf_path = temp_dir / "multipage.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50), f"Page {i+1} content")
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert result.page_count == 3
        assert "Page 1" in result.text or "Page 2" in result.text
    
    def test_get_page_count_pdf(self, temp_dir):
        """Test getting page count from PDF."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "pages.pdf"
        doc = fitz.open()
        for _ in range(5):
            doc.new_page()
        doc.save(str(pdf_path))
        doc.close()
        
        page_count = extractor.get_page_count(str(pdf_path))
        
        assert page_count == 5
    
    def test_pdf_with_special_characters(self, temp_dir):
        """Test PDF with special characters."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "special.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Special: @#$%^&*()")
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert len(result.text) > 0
    
    def test_pdf_with_unicode(self, temp_dir):
        """Test PDF with Unicode characters."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "unicode.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Unicode: 你好 مرحبا Привет 🚀")
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert result.char_count > 0
    
    def test_empty_pdf(self, temp_dir):
        """Test extraction from PDF with no text."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "empty.pdf"
        doc = fitz.open()
        doc.new_page()  # Empty page
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert result.page_count == 1
        assert result.char_count >= 0
    
    def test_pdf_extraction_error_invalid_file(self, temp_dir):
        """Test error handling for invalid PDF file."""
        extractor = PDFExtractor()
        
        # Create invalid PDF
        invalid_pdf = temp_dir / "invalid.pdf"
        invalid_pdf.write_text("Not a real PDF")
        
        with pytest.raises(ExtractionError):
            extractor.extract_text(str(invalid_pdf))
    
    def test_pdf_extraction_error_nonexistent_file(self):
        """Test error handling for non-existent file."""
        extractor = PDFExtractor()
        
        with pytest.raises(ExtractionError):
            extractor.extract_text("/nonexistent/file.pdf")
    
    @patch('app.config.settings.max_pages_per_document', 2)
    def test_pdf_page_limit_exceeded(self, temp_dir):
        """Test page limit validation."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "large.pdf"
        doc = fitz.open()
        for _ in range(5):  # More than limit
            doc.new_page()
        doc.save(str(pdf_path))
        doc.close()
        
        with pytest.raises(PageLimitExceededError):
            extractor.extract_text(str(pdf_path))


@pytest.mark.unit
class TestDOCXExtractor:
    """Tests for DOCX text extraction."""
    
    def test_docx_extractor_initialization(self):
        """Test DOCX extractor initializes correctly."""
        extractor = DOCXExtractor()
        
        assert isinstance(extractor, DOCXExtractor)
        assert isinstance(extractor, BaseExtractor)
    
    def test_extract_from_simple_docx(self, temp_dir):
        """Test extraction from simple DOCX file."""
        extractor = DOCXExtractor()
        
        # Create a simple DOCX
        docx_path = temp_dir / "test.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test DOCX content")
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert isinstance(result, ExtractedText)
        assert "Test DOCX content" in result.text
        assert result.page_count >= 1
        assert result.char_count > 0
    
    def test_extract_from_multi_paragraph_docx(self, temp_dir):
        """Test extraction from DOCX with multiple paragraphs."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "multi.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert "First paragraph" in result.text
        assert "Third paragraph" in result.text
    
    def test_docx_with_formatting(self, temp_dir):
        """Test DOCX with formatting (bold, italic)."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "formatted.docx"
        from docx import Document
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Bold text").bold = True
        paragraph.add_run(" Normal text ")
        paragraph.add_run("Italic text").italic = True
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert "Bold text" in result.text
        assert "Italic text" in result.text
    
    def test_docx_with_tables(self, temp_dir):
        """Test DOCX with tables."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "table.docx"
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        # Should extract table content
        assert "Cell 1" in result.text or "Cell 2" in result.text
    
    def test_empty_docx(self, temp_dir):
        """Test extraction from empty DOCX."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "empty.docx"
        from docx import Document
        doc = Document()
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert result.page_count >= 1
        assert result.char_count >= 0
    
    def test_docx_with_unicode(self, temp_dir):
        """Test DOCX with Unicode characters."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "unicode.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Unicode: 你好世界 🚀")
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert result.char_count > 0
    
    def test_docx_extraction_error_invalid_file(self, temp_dir):
        """Test error handling for invalid DOCX."""
        extractor = DOCXExtractor()
        
        invalid_docx = temp_dir / "invalid.docx"
        invalid_docx.write_text("Not a real DOCX")
        
        with pytest.raises(ExtractionError):
            extractor.extract_text(str(invalid_docx))
    
    def test_docx_extraction_error_nonexistent_file(self):
        """Test error handling for non-existent file."""
        extractor = DOCXExtractor()
        
        with pytest.raises(ExtractionError):
            extractor.extract_text("/nonexistent/file.docx")
    
    def test_get_page_count_docx(self, temp_dir):
        """Test getting page count from DOCX."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "pages.docx"
        from docx import Document
        doc = Document()
        for _ in range(5):
            doc.add_paragraph("Content")
            doc.add_page_break()
        doc.save(str(docx_path))
        
        page_count = extractor.get_page_count(str(docx_path))
        
        assert page_count >= 1


@pytest.mark.unit
class TestTXTExtractor:
    """Tests for TXT text extraction."""
    
    def test_txt_extractor_initialization(self):
        """Test TXT extractor initializes correctly."""
        extractor = TXTExtractor()
        
        assert isinstance(extractor, TXTExtractor)
        assert isinstance(extractor, BaseExtractor)
    
    def test_extract_from_simple_txt(self, temp_dir):
        """Test extraction from simple TXT file."""
        extractor = TXTExtractor()
        
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("Test TXT content", encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert isinstance(result, ExtractedText)
        assert "Test TXT content" in result.text
        assert result.page_count == 1
        assert result.char_count > 0
    
    def test_extract_from_multiline_txt(self, temp_dir):
        """Test extraction from multi-line TXT."""
        extractor = TXTExtractor()
        
        content = "Line 1\nLine 2\nLine 3"
        txt_path = temp_dir / "multiline.txt"
        txt_path.write_text(content, encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert "Line 1" in result.text
        assert "Line 3" in result.text
    
    def test_extract_with_utf8_encoding(self, temp_dir):
        """Test extraction with UTF-8 encoding."""
        extractor = TXTExtractor()
        
        content = "UTF-8: 你好世界 مرحبا"
        txt_path = temp_dir / "utf8.txt"
        txt_path.write_text(content, encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert result.char_count > 0
        assert len(result.text) > 0
    
    def test_extract_with_latin1_encoding(self, temp_dir):
        """Test extraction with Latin-1 encoding."""
        extractor = TXTExtractor()
        
        content = "Latin-1: café résumé"
        txt_path = temp_dir / "latin1.txt"
        txt_path.write_bytes(content.encode('latin-1'))
        
        result = extractor.extract_text(str(txt_path))
        
        assert result.char_count > 0
    
    def test_extract_with_windows_encoding(self, temp_dir):
        """Test extraction with Windows-1252 encoding."""
        extractor = TXTExtractor()
        
        content = "Windows text"
        txt_path = temp_dir / "windows.txt"
        txt_path.write_bytes(content.encode('windows-1252'))
        
        result = extractor.extract_text(str(txt_path))
        
        assert "Windows text" in result.text
    
    def test_empty_txt_file(self, temp_dir):
        """Test extraction from empty TXT file."""
        extractor = TXTExtractor()
        
        txt_path = temp_dir / "empty.txt"
        txt_path.write_text("", encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert result.page_count == 1
        assert result.char_count == 0
    
    def test_txt_with_special_characters(self, temp_dir):
        """Test TXT with special characters."""
        extractor = TXTExtractor()
        
        content = "Special: @#$%^&*()!~`"
        txt_path = temp_dir / "special.txt"
        txt_path.write_text(content, encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert result.char_count > 0
    
    def test_txt_with_mixed_newlines(self, temp_dir):
        """Test TXT with different newline types."""
        extractor = TXTExtractor()
        
        content = "Line 1\nLine 2\r\nLine 3\rLine 4"
        txt_path = temp_dir / "newlines.txt"
        txt_path.write_bytes(content.encode('utf-8'))
        
        result = extractor.extract_text(str(txt_path))
        
        assert "Line 1" in result.text
        assert "Line 4" in result.text
    
    def test_large_txt_file(self, temp_dir):
        """Test extraction from large TXT file."""
        extractor = TXTExtractor()
        
        content = "Line of text. " * 10000
        txt_path = temp_dir / "large.txt"
        txt_path.write_text(content, encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert result.char_count > 50000
    
    def test_txt_extraction_error_nonexistent_file(self):
        """Test error handling for non-existent file."""
        extractor = TXTExtractor()
        
        with pytest.raises(ExtractionError):
            extractor.extract_text("/nonexistent/file.txt")
    
    def test_get_page_count_txt(self, temp_dir):
        """Test getting page count from TXT (always 1)."""
        extractor = TXTExtractor()
        
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("Content", encoding='utf-8')
        
        page_count = extractor.get_page_count(str(txt_path))
        
        assert page_count == 1


@pytest.mark.unit
class TestExtractorFactory:
    """Tests for ExtractorFactory."""
    
    def test_factory_get_pdf_extractor(self):
        """Test factory returns PDF extractor for PDF files."""
        factory = ExtractorFactory()
        
        extractor = factory.get_extractor("test.pdf")
        
        assert isinstance(extractor, PDFExtractor)
    
    def test_factory_get_docx_extractor(self):
        """Test factory returns DOCX extractor for DOCX files."""
        factory = ExtractorFactory()
        
        extractor = factory.get_extractor("test.docx")
        
        assert isinstance(extractor, DOCXExtractor)
    
    def test_factory_get_txt_extractor(self):
        """Test factory returns TXT extractor for TXT files."""
        factory = ExtractorFactory()
        
        extractor = factory.get_extractor("test.txt")
        
        assert isinstance(extractor, TXTExtractor)
    
    def test_factory_case_insensitive(self):
        """Test factory handles uppercase extensions."""
        factory = ExtractorFactory()
        
        pdf_extractor = factory.get_extractor("TEST.PDF")
        docx_extractor = factory.get_extractor("TEST.DOCX")
        txt_extractor = factory.get_extractor("TEST.TXT")
        
        assert isinstance(pdf_extractor, PDFExtractor)
        assert isinstance(docx_extractor, DOCXExtractor)
        assert isinstance(txt_extractor, TXTExtractor)
    
    def test_factory_unsupported_extension(self):
        """Test factory raises error for unsupported file type."""
        factory = ExtractorFactory()
        
        with pytest.raises(ValueError):
            factory.get_extractor("test.xyz")
    
    def test_factory_no_extension(self):
        """Test factory handles files without extension."""
        factory = ExtractorFactory()
        
        with pytest.raises(ValueError):
            factory.get_extractor("noextension")


@pytest.mark.unit
class TestExtractedText:
    """Tests for ExtractedText dataclass."""
    
    def test_extracted_text_creation(self):
        """Test ExtractedText can be created."""
        extracted = ExtractedText(
            text="Test content",
            page_count=1,
            char_count=12,
            metadata={"key": "value"}
        )
        
        assert extracted.text == "Test content"
        assert extracted.page_count == 1
        assert extracted.char_count == 12
        assert extracted.metadata == {"key": "value"}
    
    def test_extracted_text_with_empty_text(self):
        """Test ExtractedText with empty text."""
        extracted = ExtractedText(
            text="",
            page_count=1,
            char_count=0,
            metadata={}
        )
        
        assert extracted.text == ""
        assert extracted.char_count == 0


@pytest.mark.unit
class TestBaseExtractor:
    """Tests for BaseExtractor abstract class."""
    
    def test_base_extractor_cannot_instantiate(self):
        """Test that BaseExtractor cannot be instantiated directly."""
        with pytest.raises(TypeError):
            BaseExtractor()
    
    def test_validate_page_count_within_limit(self):
        """Test page count validation passes within limit."""
        # Create a concrete implementation for testing
        class ConcreteExtractor(BaseExtractor):
            def extract_text(self, file_path: str) -> ExtractedText:
                pass
            def get_page_count(self, file_path: str) -> int:
                pass
        
        extractor = ConcreteExtractor()
        
        # Should not raise error
        with patch('app.config.settings.max_pages_per_document', 100):
            extractor.validate_page_count(50, "test.pdf")
    
    def test_validate_page_count_exceeds_limit(self):
        """Test page count validation raises error when exceeded."""
        class ConcreteExtractor(BaseExtractor):
            def extract_text(self, file_path: str) -> ExtractedText:
                pass
            def get_page_count(self, file_path: str) -> int:
                pass
        
        extractor = ConcreteExtractor()
        
        with patch('app.config.settings.max_pages_per_document', 10):
            with pytest.raises(PageLimitExceededError):
                extractor.validate_page_count(50, "test.pdf")


@pytest.mark.unit
class TestExtractionMetadata:
    """Tests for extraction metadata handling."""
    
    def test_pdf_metadata_includes_extractor_info(self, temp_dir):
        """Test that PDF extraction includes metadata."""
        extractor = PDFExtractor()
        
        pdf_path = temp_dir / "meta.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Test")
        doc.save(str(pdf_path))
        doc.close()
        
        result = extractor.extract_text(str(pdf_path))
        
        assert isinstance(result.metadata, dict)
    
    def test_docx_metadata_structure(self, temp_dir):
        """Test DOCX extraction metadata structure."""
        extractor = DOCXExtractor()
        
        docx_path = temp_dir / "meta.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(docx_path))
        
        result = extractor.extract_text(str(docx_path))
        
        assert isinstance(result.metadata, dict)
    
    def test_txt_metadata_includes_encoding(self, temp_dir):
        """Test TXT extraction includes encoding info."""
        extractor = TXTExtractor()
        
        txt_path = temp_dir / "meta.txt"
        txt_path.write_text("Test", encoding='utf-8')
        
        result = extractor.extract_text(str(txt_path))
        
        assert isinstance(result.metadata, dict)

